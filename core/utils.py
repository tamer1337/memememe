# app/utils.py
from docx import Document
import html

def docx_to_html(file_path):
    doc = Document(file_path)
    html_content = []
    
    for para in doc.paragraphs:
        # Обработка текста с форматированием внутри параграфа
        text_parts = []
        for run in para.runs:
            run_text = html.escape(run.text)
            if run.bold:
                run_text = f"<strong>{run_text}</strong>"
            if run.italic:
                run_text = f"<em>{run_text}</em>"
            if run.underline:
                run_text = f"<u>{run_text}</u>"
            text_parts.append(run_text)
        
        full_text = ''.join(text_parts)
        
        # Определение типа элемента
        if para.style.name.startswith('Heading'):
            level = min(int(para.style.name[-1]), 6)
            html_content.append(f"<h{level}>{full_text}</h{level}>")
        elif para.style.name == 'List Paragraph':
            html_content.append(f"<li>{full_text}</li>")
        else:
            html_content.append(f"<p>{full_text}</p>")
    
    # Обработка таблиц
    for table in doc.tables:
        html_content.append('<table class="table-auto border-collapse w-full mb-4">')
        for row in table.rows:
            html_content.append('<tr>')
            for cell in row.cells:
                html_content.append(f'<td class="border px-4 py-2">{html.escape(cell.text)}</td>')
            html_content.append('</tr>')
        html_content.append('</table>')
    
    # Объединяем списки <li> в <ul> или <ol>
    final_html = []
    in_list = False
    
    for line in html_content:
        if line.startswith('<li>'):
            if not in_list:
                final_html.append('<ul class="list-disc pl-5 mb-4">')
                in_list = True
            final_html.append(line)
        else:
            if in_list:
                final_html.append('</ul>')
                in_list = False
            final_html.append(line)
    
    if in_list:
        final_html.append('</ul>')
    
    return '\n'.join(final_html)